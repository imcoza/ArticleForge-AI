"""
Document service for creating Word documents.
Handles document generation with text and images.
"""
import io
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import requests
from src.config.settings import settings
from src.utils.logger import logger
from src.services.image_service import image_service


class DocumentService:
    """Service for managing document operations."""
    
    @staticmethod
    def create_word_document(
        title: str,
        content: str,
        image_url: Optional[str] = None,
        metadata: Optional[dict] = None
    ) -> Document:
        """
        Create a Word document with title, content, and optional image.
        
        Args:
            title: Document title
            content: Document content
            image_url: Optional image URL to include
            metadata: Optional metadata dictionary
            
        Returns:
            Document object
        """
        try:
            logger.info(f"Creating Word document: {title}")
            
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata if provided (only word count, no date)
            if metadata and metadata.get('word_count'):
                doc.add_paragraph(f"Word count: {metadata.get('word_count')}")
            
            # Add content
            doc.add_paragraph(content)
            
            # Add image if provided
            if image_url:
                try:
                    image_data = image_service.download_image(image_url)
                    if image_data:
                        image_stream = io.BytesIO(image_data)
                        image = Image.open(image_stream)
                        
                        # Add image heading
                        doc.add_heading("Reference Image", level=2)
                        
                        # Add image to document
                        image_stream.seek(0)
                        doc.add_picture(image_stream, width=Inches(settings.IMAGE_WIDTH_INCHES))
                        
                        logger.info("Image added to document successfully")
                except Exception as e:
                    logger.warning(f"Could not add image to document: {str(e)}")
            
            logger.info("Word document created successfully")
            return doc
            
        except Exception as e:
            logger.error(f"Error creating document: {str(e)}", exc_info=True)
            raise
    
    @staticmethod
    def save_document_to_bytes(doc: Document) -> bytes:
        """
        Save document to bytes buffer.
        
        Args:
            doc: Document object
            
        Returns:
            Document as bytes
        """
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()




